from docx import Document
import os, re

class FormatterAgent:
    name = "formatter"

    def __init__(self, template_path):
        """
        template_path: ruta ABSOLUTA enviada por CoordinatorAgent
        """

        if not template_path or not os.path.exists(template_path):
            raise FileNotFoundError(f"[FormatterAgent] NO se encontró la plantilla → {template_path}")

        self.template_path = template_path
        print(f"[FormatterAgent] Usando plantilla REAL: {self.template_path}")

    @staticmethod
    def _replace_in_paragraph(paragraph, data):
        full_text = "".join(run.text for run in paragraph.runs)

        for key, value in data.items():
            pattern = r"\{\s*" + re.escape(key) + r"\s*\}"
            full_text = re.sub(pattern, str(value), full_text)

        remaining = full_text
        for run in paragraph.runs:
            n = len(run.text)
            run.text = remaining[:n]
            remaining = remaining[n:]

        if remaining:
            paragraph.runs[-1].text += remaining

    def _replace_in_table(self, table, data):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._replace_in_paragraph(paragraph, data)

    def generate_contract(self, data: dict, output_path: str):
        doc = Document(self.template_path)

        for p in doc.paragraphs:
            self._replace_in_paragraph(p, data)

        for t in doc.tables:
            self._replace_in_table(t, data)

        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        doc.save(output_path)
        print(f"[FormatterAgent] Contrato generado en: {output_path}")

